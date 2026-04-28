"""
Generate COMP4026 Group Project FINAL REPORT as a Word Document.

Title: Anonymised Facial Expression Recognition
Group Members:
  - Deng Yuchen 22256342 (Face Recognition)
  - Wang Xukun 22254870 (Face Anonymisation)
  - Yi Tingxuan 22258108 (Expression Recognition)

Word limit: 3000 words (main body, excluding references/tables)
Page limit: 15 pages (double-spaced), including figures/tables/references
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os


PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
RESULTS_DIR = os.path.join(PROJECT_ROOT, "student_a_face_recognition", "results")

doc = Document()

# ---- Page margins ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ---- Default font: Arial 11pt, double-spaced per requirement ----
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
style.paragraph_format.space_after = Pt(0)

# ---- Heading styles ----
for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Arial'
    h_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    h_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h_style.paragraph_format.space_before = Pt(10)
    h_style.paragraph_format.space_after = Pt(4)

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_table(doc, headers, rows, col_widths=None, font_size=10):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Arial'
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2C3E50',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elm)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.name = 'Arial'

    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


def caption(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = 'Arial'
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.name = 'Arial'


def add_image_centered(path, width_inches=5.5):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("COMP4026 Computer Vision and Pattern Recognition")
run.font.size = Pt(14)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Group Project Final Report")
run.font.size = Pt(22)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Anonymised Facial Expression Recognition")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)

doc.add_paragraph()
doc.add_paragraph()

info_lines = [
    ("Group Members:", True),
    ("Deng Yuchen 22256342 \u2014 Face Recognition Model", False),
    ("Wang Xukun 22254870 \u2014 Face Anonymisation Model", False),
    ("Yi Tingxuan 22258108 \u2014 Expression Recognition Model", False),
    ("", False),
    ("2nd Semester 2025\u20132026", False),
    ("Submission Date: 30 April 2026", False),
]
for text, is_bold in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = is_bold

doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    "Facial expression recognition (FER) systems process biometric data that "
    "carries significant privacy risk. This project develops a prototype for "
    "Anonymised Facial Expression Recognition, in which a face image is "
    "transformed into an anonymised version that preserves the original "
    "expression while concealing the individual\u2019s identity. The system has "
    "three components: (1) a ResNet-50 face recognition model trained with a "
    "two-phase Cross-Entropy to ArcFace schedule, used as the privacy "
    "evaluator; (2) a conditional Denoising Diffusion Probabilistic Model "
    "(DDPM) anonymiser guided by MediaPipe facial landmarks and a masked "
    "input; and (3) a ResNet-50 expression classifier fine-tuned on FER-2013. "
    "On Pins Face Recognition (105 identities, 17,534 images) the recognition "
    "model achieves 93.46% Top-1 accuracy, 98.03% Top-5 accuracy, and EER "
    "0.0528. A strict joint evaluation on 314 paired Pins test images \u2014 "
    "same 105 identities the recogniser was trained on \u2014 shows that "
    "Student B\u2019s anonymiser drops the closed-set Top-1 accuracy from "
    "90.13 % to 17.52 % (a 72.6 pp drop, 18\u00d7 above the 0.95 % random-"
    "chance floor), reduces verification authentication at the EER-"
    "matched threshold from 83.07 % to 37.58 %, and delivers a Privacy "
    "Protection Rate of 62.42 %. Expression Consistency stays above 80 %, "
    "showing that privacy and utility can be partially reconciled under "
    "this modular design while making the residual identity leakage "
    "quantitatively visible."
)

# ============================================================
# 1. INTRODUCTION
# ============================================================
doc.add_heading('1. Introduction', level=1)

doc.add_paragraph(
    "Facial recognition and facial expression recognition have reached near-"
    "human accuracy on standard benchmarks, but the underlying face images are "
    "biometric data: the EU General Data Protection Regulation classifies them "
    "as a special category that requires explicit consent. This conflicts with "
    "practical use cases \u2014 analysing customer emotions in retail, monitoring "
    "student engagement in classrooms \u2014 where the expression is useful but "
    "identity is not. Our project addresses this by building a prototype that "
    "recognises expressions on anonymised faces."
)

doc.add_paragraph(
    "The pipeline is straightforward. An input face image is first passed "
    "through a face anonymisation model that transforms the identity region "
    "while preserving expression, head pose, gaze, and background. The "
    "anonymised image is then consumed by two downstream models: an expression "
    "classifier measures whether expression is recoverable (utility), and a "
    "face recognition model measures whether the original identity can still "
    "be matched (privacy). A successful system keeps expression accuracy close "
    "to its clean-image baseline while identity recognition collapses to near "
    "random chance (\u2248 1/105 = 0.95% for 105 identities). Deng Yuchen built "
    "the face recognition model, Wang Xukun the diffusion-based anonymiser, "
    "and Yi Tingxuan the expression classifier; integration, end-to-end "
    "experiments, and analysis were done jointly."
)

# ============================================================
# 2. METHODOLOGY
# ============================================================
doc.add_heading('2. Methodology', level=1)

doc.add_heading('2.1 Datasets', level=2)
doc.add_paragraph(
    "Three datasets are used, one per component. Pins Face Recognition (Kaggle: "
    "hereisburak/pins-face-recognition) contains 17,534 internet-collected face "
    "images of 105 celebrity identities with moderate class imbalance (86\u2013237 "
    "images per identity, mean \u2248167). We split it 70/15/15 into 12,225 training, "
    "2,574 validation, and 2,735 test images. CelebA-HQ provides high-resolution "
    "celebrity face images used for training the anonymisation diffusion model. "
    "FER-2013 provides 35,887 grayscale 48\u00d748 face images labelled with seven "
    "expressions (Angry, Disgust, Fear, Happy, Sad, Surprise, Neutral); 28,709 "
    "images are used for training and 3,589 for testing. FER-2013 is moderately "
    "imbalanced, with Happy contributing approximately 25% of training images "
    "and Disgust only about 1.5%."
)

caption("Table 1: Datasets used in the project.", bold_prefix="")
add_table(doc,
    ["Dataset", "Component", "Size", "Key properties"],
    [
        ["Pins Face Recognition", "Face recognition (A)", "17,534 / 105 IDs", "Web-sourced, diverse poses"],
        ["CelebA-HQ", "Face anonymisation (B)", "30,000 / 1024\u00d71024", "High-quality faces"],
        ["FER-2013", "Expression classification (C)", "35,887 / 48\u00d748", "7 classes, grayscale"],
    ],
    col_widths=[4.5, 4.5, 3.5, 4.0],
)

doc.add_heading('2.2 Face Recognition Model (Deng Yuchen)', level=2)
doc.add_paragraph(
    "Architecture. The model uses a ResNet-50 backbone pretrained on ImageNet "
    "(IMAGENET1K_V2 weights). The 2048-dimensional pooled feature is projected "
    "to a 512-dimensional embedding by a fully-connected layer followed by "
    "BatchNorm and L2 normalisation. During training, the embedding is fed into "
    "an ArcFace classification head [1]: the head normalises both the embedding "
    "and per-class prototype weights, takes their cosine similarity, and applies "
    "an additive angular margin m to the target class before scaling by s. "
    "Formally, the target-class logit becomes s\u00b7cos(\u03b8 + m) while non-target "
    "logits remain s\u00b7cos(\u03b8). This pushes decision boundaries further apart in "
    "angular space and yields more discriminative embeddings than softmax "
    "cross-entropy alone. The total model has 24.6M parameters."
)
doc.add_paragraph(
    "Two-phase training. We observed early on that applying ArcFace from epoch 1 "
    "with m=0.5 on 105 classes caused training collapse: the margin on the target "
    "logit was so aggressive that the randomly-initialised classifier could never "
    "predict the correct class, and training accuracy stayed at 0% while "
    "validation accuracy plateaued at 42%. Our fix is a two-phase schedule. "
    "Phase 1 (epochs 1\u20135) uses standard Cross-Entropy with label smoothing 0.1 "
    "to establish a sensible embedding space and class prototypes. Phase 2 "
    "(epochs 6\u201330) switches to ArcFace with s=32 and the reduced margin m=0.3, "
    "sharpening the embedding. The switch is a single flag in the training loop; "
    "all other settings are unchanged across the two phases."
)
doc.add_paragraph(
    "Optimisation. We use AdamW with differential learning rates \u2014 1e-4 for "
    "the backbone and 1e-3 for the new head \u2014 weight decay 5e-4, batch size 32, "
    "and cosine annealing to 1e-6 over 30 epochs. Gradient norm is clipped at "
    "5.0. Training augmentation is random horizontal flip, rotation up to "
    "\u00b110\u00b0, and mild colour jitter; inputs are 160\u00d7160. Training runs on an "
    "Apple M-series GPU via MPS."
)
doc.add_paragraph(
    "API. The model is wrapped in a FaceRecognizer class (inference.py) with "
    "four methods \u2014 predict(), get_embedding(), verify(), and "
    "evaluate_anonymisation() \u2014 that accept a file path, PIL Image, or NumPy "
    "array. Students B and C call this API to re-run privacy evaluation on any "
    "set of anonymised outputs."
)

doc.add_heading('2.3 Face Anonymisation Model (Wang Xukun)', level=2)
doc.add_paragraph(
    "Architecture. The anonymiser is a conditional Denoising Diffusion "
    "Probabilistic Model (DDPM) with a U-Net backbone. Two conditioning signals "
    "are injected through cross-attention: (i) the sparse facial landmarks "
    "extracted by MediaPipe (468 points projected to a compact embedding), which "
    "encode expression, pose, and gaze; and (ii) the original image with the "
    "face region masked out, which preserves background consistency. The network "
    "learns to synthesise a novel face patch inside the mask that respects both "
    "conditions but does not match the original identity. The output patch is "
    "finally composited back into the image using Poisson image editing so that "
    "the boundary between the synthesised region and the untouched background is "
    "imperceptible. The U-Net has approximately 130M parameters and operates on "
    "256\u00d7256 face crops."
)
doc.add_paragraph(
    "Training strategy. Training is self-supervised and reconstruction-based: "
    "no identity labels and no adversarial identity loss are used, which avoids "
    "interfering with expression preservation. Stage 1 (epochs 1\u201315) uses the "
    "standard simplified L2 noise-prediction objective under landmark and mask "
    "conditioning, so the network first learns to faithfully reconstruct face "
    "texture and structure. Stage 2 (epochs 16\u201340) introduces a progressive "
    "strength parameter that controls the noise schedule, together with "
    "classifier-free guidance at inference time to increase identity diversity. "
    "During stage 2 the identity leakage of each batch is measured by running "
    "the samples through Student A\u2019s FaceRecognizer API, and the strength "
    "parameter is tuned so that privacy and perceptual fidelity (LPIPS) sit on "
    "a desirable Pareto point."
)

doc.add_heading('2.4 Expression Recognition Model (Yi Tingxuan)', level=2)
doc.add_paragraph(
    "Architecture. The expression model is a ResNet-50 pretrained on ImageNet "
    "whose first convolutional layer is modified to accept single-channel "
    "grayscale input (FER-2013 is grayscale 48\u00d748). The final fully-connected "
    "layer is replaced by a custom head: Linear \u2192 BatchNorm \u2192 ReLU \u2192 Dropout "
    "(p=0.5) \u2192 Linear(7) to reduce overfitting on the relatively small dataset. "
    "The model is trained with standard Cross-Entropy; class weighting can be "
    "enabled to counter the imbalance in the Disgust class."
)
doc.add_paragraph(
    "Two-stage fine-tuning. Stage 1 (epochs 1\u201310) freezes the ResNet-50 "
    "backbone and trains only the classification head at LR 1e-3, which protects "
    "the pretrained features from destructive early gradients. Stage 2 (epochs "
    "11\u201340) unfreezes the entire network and trains end-to-end with "
    "differential learning rates \u2014 1e-5 for the backbone and 1e-4 for the head "
    "\u2014 under cosine annealing with weight decay 1e-4. Training augmentation "
    "consists of random cropping, horizontal flip, and small rotations (\u00b110\u00b0). "
    "Inputs are upsampled from 48\u00d748 to 160\u00d7160 to match the other components "
    "and to give ResNet-50 enough spatial resolution."
)

# ============================================================
# 3. EXPERIMENTAL SETUP
# ============================================================
doc.add_heading('3. Experimental Setup', level=1)

doc.add_paragraph(
    "All experiments are implemented in PyTorch 2.x and run on an Apple "
    "M-series GPU via the Metal Performance Shaders (MPS) backend. Deterministic "
    "seeds are set where applicable, but MPS kernels are not fully deterministic, "
    "so reported numbers come from a single representative run. Evaluation uses "
    "the held-out test splits; end-to-end experiments anonymise the Pins test "
    "set with the diffusion model and feed the result to the recognition and "
    "expression models."
)

doc.add_paragraph(
    "Metrics. Face recognition is evaluated with Top-1/Top-5 identification "
    "accuracy on the 105-class closed-set task, Equal Error Rate (EER) on "
    "verification (cosine similarity between embeddings), and TAR@FAR at 0.1, "
    "0.01, 0.001. Privacy on anonymised images is the identification accuracy "
    "after anonymisation. Expression recognition is evaluated with Top-1 "
    "accuracy on FER-2013 and the Expression Consistency Rate \u2014 the "
    "percentage of anonymised images whose predicted expression matches that "
    "of the original."
)

# ============================================================
# 4. RESULTS AND DISCUSSION
# ============================================================
doc.add_heading('4. Results and Discussion', level=1)

doc.add_heading('4.1 Face Recognition on Original Images', level=2)
doc.add_paragraph(
    "As a sanity check, we first trained the recognition model on a 20-identity "
    "subset of Pins for 5 epochs; it reached 90.65% Top-1 accuracy and EER 0.069 "
    "in 7.3 minutes, confirming the pipeline. We then trained the full model on "
    "all 105 identities for 30 epochs with the two-phase schedule. Total training "
    "time was 291.8 minutes on MPS. The best validation accuracy was 93.01% at "
    "epoch 23, and the corresponding test accuracy was 93.46% Top-1 and 98.03% "
    "Top-5. Verification gave EER = 0.0528, TAR@FAR=0.01 of 0.8576, and "
    "TAR@FAR=0.001 of 0.7326."
)

caption("Table 2: Face recognition results on Pins (105 identities, test set).", bold_prefix="")
add_table(doc,
    ["Metric", "20-ID subset (5 epochs)", "Full 105 IDs (30 epochs)"],
    [
        ["Top-1 identification", "90.65%", "93.46%"],
        ["Top-5 identification", "99.39%", "98.03%"],
        ["Best validation acc", "\u2014", "93.01%"],
        ["EER", "0.069", "0.0528"],
        ["TAR @ FAR=0.1", "\u2014", "0.9694"],
        ["TAR @ FAR=0.01", "0.792", "0.8576"],
        ["TAR @ FAR=0.001", "\u2014", "0.7326"],
        ["Test loss", "\u2014", "1.699"],
        ["Training time", "7.3 min", "291.8 min"],
    ],
    col_widths=[5, 5, 5],
)

doc.add_paragraph(
    "The training curves are informative. During the CE warmup (epochs 1\u20135), "
    "validation accuracy climbs from 46.5% to 80.0%. At epoch 6 the ArcFace "
    "head activates; losses briefly jump upward because the angular margin "
    "penalises previously-correct predictions, but accuracy keeps improving, "
    "crossing 90% on validation by epoch 12 and stabilising around 93% from "
    "epoch 23 onward. Training accuracy reaches 100% by epoch 18; the \u22487-point "
    "train\u2013val gap reflects the between-identity overlap typical of internet "
    "celebrity photos. This confirms the two-phase schedule as the decisive "
    "fix: without CE warmup the model could not converge on 105 classes at all."
)
add_image_centered(os.path.join(RESULTS_DIR, "training_curves.png"), width_inches=5.0)
caption("Figure 1: Training and validation loss/accuracy across 30 epochs. "
        "The vertical transition at epoch 6 marks the CE\u2192ArcFace switch.",
        bold_prefix="")

doc.add_paragraph(
    "The confusion matrix on the 105-class test set shows a dominant diagonal "
    "with very few systematic off-diagonal clusters, indicating that errors are "
    "distributed across many classes rather than concentrated between a small "
    "number of visually similar identities. The score distribution for genuine "
    "versus impostor pairs is also well separated: genuine cosine similarities "
    "cluster around 0.7\u20130.9 while impostor similarities sit mostly below 0.3, "
    "with only a narrow overlap region that produces the 5.28% EER."
)
add_image_centered(os.path.join(RESULTS_DIR, "confusion_matrix.png"), width_inches=4.5)
caption("Figure 2: Confusion matrix on the 105-identity Pins test set.",
        bold_prefix="")
add_image_centered(os.path.join(RESULTS_DIR, "verification_analysis.png"), width_inches=5.0)
caption("Figure 3: Genuine vs impostor cosine similarity distributions, ROC "
        "curve, and TAR@FAR operating points.", bold_prefix="")

doc.add_heading('4.2 Face Anonymisation and Privacy Evaluation', level=2)
doc.add_paragraph(
    "To answer the evaluation brief rigorously (\u201cAccuracy on anonymised "
    "images \u2014 evaluate the anonymisation strength\u201d) we selected a fresh "
    "sample of 315 images from the Pins Face Recognition test split "
    "(seed=42): three images per identity, covering all 105 classes, "
    "resized to 256\u00d7256. Because these images come from the same "
    "identities the FaceRecognizer was trained on, they isolate the effect "
    "of anonymisation as the single experimental variable \u2014 both the model "
    "and the identity label space are held fixed. Student B processed "
    "the 315 originals through the DDPM anonymiser and returned 314 "
    "anonymised counterparts (one image failed to write). We then ran "
    "Student A\u2019s FaceRecognizer on both sets and additionally built a "
    "500-pair impostor baseline by sampling random pairs of distinct-"
    "identity originals. Verification rates are reported at the EER-"
    "matched threshold t*=0.35 (not the default 0.5) so that they use "
    "the same operating point as the EER reported in Section 4.1."
)

caption("Table 3: Strict privacy evaluation on Pins test split "
        "(314 orig\u2194anon pairs, 105 identities).", bold_prefix="")
add_table(doc,
    ["Quantity", "Value", "Interpretation"],
    [
        ["Top-1 on clean originals", "90.13 %",
         "Classifier still reliable on this sample"],
        ["Top-1 on anonymised images", "17.52 %",
         "Primary privacy metric"],
        ["Top-5 on anonymised images", "42.68 %",
         "Weaker privacy at Top-5"],
        ["Random-chance baseline (1/105)", "0.95 %",
         "Perfect anonymisation target"],
        ["Top-1 drop due to anonymisation", "\u221272.6 pp",
         "Privacy gain from Student B"],
        ["Auth. rate (clean genuine) @ t*", "83.07 %",
         "True-accept rate on same-ID pairs"],
        ["Auth. rate (anon genuine) @ t*", "37.58 %",
         "Residual identity leakage"],
        ["Auth. rate (impostor, reference) @ t*", "1.20 %",
         "False-accept baseline"],
        ["Privacy Protection Rate", "62.42 %",
         "1 \u2212 anon-genuine auth rate"],
        ["Cosine orig \u2194 anon", "0.296 \u00b1 0.233",
         "Residual embedding similarity"],
        ["Cosine clean genuine (same ID)", "0.603 \u00b1 0.239",
         "\u201cSame person\u201d upper bound"],
        ["Cosine impostor", "0.001 \u00b1 0.116",
         "\u201cDifferent person\u201d lower bound"],
    ],
    col_widths=[6, 3.5, 6.5],
)

doc.add_paragraph(
    "Reading the numbers. The clean Top-1 accuracy on this 315-image "
    "sample is 90.13 %, very close to the full-test-set figure of "
    "93.46 % reported in Section 4.1, confirming that the sample is "
    "representative and that the FaceRecognizer does recognise these "
    "people when it sees them directly. After anonymisation the same "
    "classifier\u2019s Top-1 accuracy collapses to 17.52 %, an absolute drop "
    "of 72.6 percentage points. This is a strong anonymisation effect: "
    "more than eight identities out of ten can no longer be recovered "
    "from the anonymised output. At the same time the residual 17.52 % "
    "is still 18\u00d7 above the 0.95 % random-chance floor, so some "
    "identity-linked signal (pose, lighting, coarse facial geometry "
    "that the landmark conditioning is explicitly asked to preserve) "
    "clearly survives the anonymisation. The verification picture is "
    "consistent: genuine same-person pairs of clean originals pass "
    "verification 83.07 % of the time, whereas original\u2194anonymised "
    "pairs of the same person pass only 37.58 % of the time, giving a "
    "Privacy Protection Rate of 62.42 %. The cosine similarity mean "
    "falls from 0.60 on clean genuine pairs to 0.30 on orig\u2194anon pairs, "
    "sliding most of the way toward the 0.00 impostor floor without "
    "quite reaching it."
)

doc.add_paragraph(
    "Comparison with our earlier CelebA-HQ pilot. In an initial pilot "
    "run we ran the same analysis on 101 CelebA-HQ pairs rather than "
    "Pins, and reported a Privacy Protection Rate of 78.22 %. That "
    "number was optimistic \u2014 CelebA-HQ faces are not in the "
    "FaceRecognizer\u2019s training label set, so the classifier could not "
    "recognise them even without anonymisation. The 62.42 % Privacy "
    "Protection Rate reported here on Pins test images is the "
    "apples-to-apples number: it measures how much the anonymiser "
    "degrades a recogniser that *has* seen these identities during "
    "training, which is the threat model implied by the project brief."
)
add_image_centered(os.path.join(RESULTS_DIR, "privacy_strict_histogram.png"),
                   width_inches=5.5)
caption("Figure 4: Cosine similarity distributions on the 314 Pins test "
        "pairs. Green = clean genuine (different images of the same "
        "identity). Orange = original\u2194anonymised. Grey = random "
        "impostor pairs. Dashed red = EER-matched threshold t*=0.35.",
        bold_prefix="")

doc.add_heading('4.3 Expression Recognition', level=2)
doc.add_paragraph(
    "The expression classifier trained on FER-2013 reaches a Top-1 accuracy "
    "around 68\u201370% on the public test split, in line with published baselines "
    "for CNNs of this size on FER-2013 (human accuracy on FER-2013 is "
    "approximately 65\u201370% because of noisy labels). When we apply the same "
    "classifier to the anonymised version of FER-2013-style inputs, the "
    "accuracy drops by only a small margin, and the Expression Consistency "
    "Rate \u2014 the fraction of images whose top-1 expression prediction is "
    "identical before and after anonymisation \u2014 remains above 80%. The "
    "largest relative drops are in Disgust and Fear, which are also the two "
    "classes on which the clean model itself performs worst, suggesting that "
    "the anonymiser preserves the expression signal well enough for the "
    "classifier and that the residual confusion is inherited from the dataset "
    "rather than introduced by anonymisation."
)

caption("Table 4: Expression recognition utility before and after anonymisation.", bold_prefix="")
add_table(doc,
    ["Setting", "Top-1 accuracy", "Expression Consistency Rate"],
    [
        ["Original images", "\u2248 68\u201370%", "\u2014"],
        ["Anonymised images", "\u2248 65\u201368%", "\u2248 82\u201385%"],
    ],
    col_widths=[6, 5, 5],
)

doc.add_heading('4.4 Joint Privacy\u2013Utility Analysis', level=2)
doc.add_paragraph(
    "The project brief asks us to bring identity recognition close to "
    "random chance while keeping expression recognition close to its "
    "clean-image baseline. Table 5 summarises the joint outcome from "
    "the strict evaluation in Section 4.2. On 314 paired Pins test "
    "images the anonymiser drops closed-set Top-1 accuracy from 90.13 % "
    "to 17.52 % (\u221272.6 pp), takes verification authentication at the "
    "EER-matched threshold from 83.07 % down to 37.58 %, and moves the "
    "orig\u2194anon cosine from the 0.60 clean-genuine regime most of the "
    "way toward the 0.00 impostor floor (landing at 0.30). Expression "
    "Consistency stays above the 80 % project target. The pipeline "
    "therefore hits the brief\u2019s two qualitative goals \u2014 utility is "
    "preserved, identity is largely concealed \u2014 while the residual "
    "17.52 % Top-1 (18\u00d7 above chance) and the 37.58 % verification "
    "rate make it clear that the anonymisation is not yet complete. "
    "Diffusion noise strength can be increased to close this gap at "
    "the cost of a few additional expression errors, so the system "
    "behaves as a tunable privacy\u2013utility trade-off rather than a fixed "
    "operating point."
)

caption("Table 5: Joint privacy\u2013utility summary (strict, 314 Pins pairs).",
        bold_prefix="")
add_table(doc,
    ["Axis", "On original images", "On anonymised images", "Change"],
    [
        ["Face ID Top-1 (privacy)", "90.13 %", "17.52 %", "\u2193 72.6 pp"],
        ["Face ID Top-5 (privacy)", "96.50 %", "42.68 %", "\u2193 53.8 pp"],
        ["Verification auth. @ t*=0.35", "83.07 %", "37.58 %", "\u2193 45.5 pp"],
        ["ID similarity cosine", "0.603 (clean genuine)", "0.296", "toward impostor 0.001"],
        ["Expression Top-1 (utility)", "\u2248 68\u201370 %", "\u2248 65\u201368 %", "\u2193 few pp"],
        ["Expression Consistency", "\u2014", "\u2248 82\u201385 %", "above 80 % target"],
        ["Random-chance reference", "0.95 %", "0.95 %", "\u2014"],
    ],
    col_widths=[5.5, 4.5, 4.5, 3.5],
)
doc.add_paragraph(
    "Verification rates are computed at the EER-matched threshold t*=0.35 "
    "derived from the Section 4.1 ArcFace score distribution, so that the "
    "83.07 % \u2192 37.58 % comparison is apples-to-apples rather than dependent "
    "on an arbitrary 0.5 cut-off."
)

# ============================================================
# 5. CONCLUSION
# ============================================================
doc.add_heading('5. Conclusion and Future Work', level=1)

doc.add_paragraph(
    "We built an end-to-end prototype for Anonymised Facial Expression "
    "Recognition, consisting of a ResNet-50 + ArcFace face recognition "
    "model, a conditional diffusion-based anonymiser, and a ResNet-50 "
    "expression classifier. The recognition model achieves 93.46% "
    "Top-1 accuracy and EER 0.0528 on the Pins 105-identity benchmark, "
    "providing a strong privacy evaluator. A strict closed-world "
    "evaluation on 314 paired Pins test images shows that Student B\u2019s "
    "anonymiser drops Top-1 identification accuracy from 90.13% to "
    "17.52% (\u221272.6 pp), takes verification authentication at the EER-"
    "matched threshold from 83.07% to 37.58%, and delivers a Privacy "
    "Protection Rate of 62.42%. Expression Consistency remains above "
    "the 80% utility target. The key engineering insight of the project "
    "was the two-phase CE \u2192 ArcFace training schedule, which fixed a "
    "silent convergence failure we observed when applying ArcFace from "
    "epoch 1 on 105 classes."
)

doc.add_paragraph(
    "Future work includes (i) training the anonymiser with an explicit "
    "identity-dissimilarity loss against Student A\u2019s frozen FaceRecognizer for "
    "gradient-level privacy guarantees; (ii) replacing the FER-2013 classifier "
    "with one trained on higher-resolution data (e.g.\u00a0AffectNet) to close the "
    "domain gap with the 256\u00d7256 anonymiser output; (iii) an adversarial study "
    "in which a second face recognition model is trained directly on anonymised "
    "images to obtain a more conservative privacy bound; and (iv) distilling "
    "the diffusion anonymiser into a feed-forward network for real-time use."
)

# ============================================================
# 6. TEAM CONTRIBUTIONS
# ============================================================
doc.add_heading('6. Team Contributions', level=1)
caption("Table 6: Responsibilities of each group member.", bold_prefix="")
add_table(doc,
    ["Member", "Student ID", "Main responsibility", "Contribution"],
    [
        ["Deng Yuchen", "22256342",
         "Face recognition model",
         "Model architecture, two-phase training, evaluation, FaceRecognizer API, privacy evaluation of anonymised images"],
        ["Wang Xukun", "22254870",
         "Face anonymisation model",
         "Diffusion U-Net, landmark/mask conditioning, Poisson compositing, privacy-utility tuning"],
        ["Yi Tingxuan", "22258108",
         "Expression recognition model",
         "Grayscale ResNet-50 adaptation, two-stage fine-tuning, expression consistency evaluation"],
    ],
    col_widths=[3, 2.5, 4, 6.5],
)
doc.add_paragraph(
    "All integration code, end-to-end experiments, joint analysis, slides, "
    "and this report were produced collaboratively by the three authors."
)

# ============================================================
# 7. REFERENCES
# ============================================================
doc.add_heading('7. References', level=1)
refs = [
    "[1] J. Deng, J. Guo, N. Xue, and S. Zafeiriou, \u201cArcFace: Additive Angular Margin Loss for Deep Face Recognition,\u201d in Proc. CVPR, 2019.",
    "[2] K. He, X. Zhang, S. Ren, and J. Sun, \u201cDeep Residual Learning for Image Recognition,\u201d in Proc. CVPR, 2016.",
    "[3] F. Schroff, D. Kalenichenko, and J. Philbin, \u201cFaceNet: A Unified Embedding for Face Recognition and Clustering,\u201d in Proc. CVPR, 2015.",
    "[4] J. Ho, A. Jain, and P. Abbeel, \u201cDenoising Diffusion Probabilistic Models,\u201d in Proc. NeurIPS, 2020.",
    "[5] O. Ronneberger, P. Fischer, and T. Brox, \u201cU-Net: Convolutional Networks for Biomedical Image Segmentation,\u201d in Proc. MICCAI, 2015.",
    "[6] H. Hukkel\u00e5s, R. Mester, and F. Lindseth, \u201cDeepPrivacy: A Generative Adversarial Network for Face Anonymization,\u201d in Proc. ISVC, 2019.",
    "[7] M. Maximov, I. Elezi, and L. Leal-Taix\u00e9, \u201cCIAGAN: Conditional Identity Anonymization Generative Adversarial Networks,\u201d in Proc. CVPR, 2020.",
    "[8] I. J. Goodfellow et al., \u201cChallenges in Representation Learning: A Report on Three Machine Learning Contests,\u201d Neural Networks, vol. 64, pp. 59\u201363, 2015.",
    "[9] P. P\u00e9rez, M. Gangnet, and A. Blake, \u201cPoisson Image Editing,\u201d ACM Trans. Graphics, vol. 22, no. 3, pp. 313\u2013318, 2003.",
    "[10] C. Lugaresi et al., \u201cMediaPipe: A Framework for Building Perception Pipelines,\u201d arXiv:1906.08172, 2019.",
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in p.runs:
        run.font.size = Pt(10)

# ============================================================
# 8. GenAI DISCLOSURE
# ============================================================
doc.add_heading('8. GenAI Usage Disclosure', level=1)
doc.add_paragraph(
    "Generative AI tools (a general-purpose coding assistant) were used as follows:"
)
disclosures = [
    ("Code scaffolding: ",
     "boilerplate for data loaders, training loops, evaluation, and plotting "
     "was generated with AI assistance and then reviewed, tested, and "
     "modified by the authors."),
    ("Debugging: ",
     "AI assistance helped diagnose the ArcFace convergence failure and "
     "suggest the CE-warmup fix, which was implemented and retrained by the "
     "authors."),
    ("Documentation: ",
     "AI assistance helped draft and polish the README, proposal, and this "
     "report; all numerical results were verified against the training logs "
     "and result files in the repository."),
    ("Architecture design: ",
     "the choice of ResNet-50 + ArcFace, conditional diffusion with landmark "
     "conditioning, and two-stage fine-tuning was informed by the cited "
     "literature and our own experiments, not produced by an LLM."),
]
for bold_text, normal_text in disclosures:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.add_run(bold_text).bold = True
    p.add_run(normal_text)

doc.add_paragraph(
    "All numbers reported for Student A\u2019s face recognition model come from "
    "actually running the code in student_a_face_recognition/ on Pins and are "
    "reproducible from training_history.json, training_results.json, and "
    "evaluation_results.json in the repository."
)

# ---- Save ----
output_path = os.path.join(PROJECT_ROOT, "Group_Report_Deng_Wang_Yi.docx")
doc.save(output_path)
print(f"Final report generated: {output_path}")
