"""
Generate COMP4026 Written Assignment as a Word Document, then convert to PDF.

Student ID: 22256342
Course: COMP4026 Computer Vision and Pattern Recognition (2025-26 S2)
Due: 30 April 2026

Part (i):  Selected YOLO algorithm (post-2018) -- YOLOv7 (CVPR 2023)
Part (ii): Essay on two unmet needs of SOTA visual object detection
           (under 400 words)

Output: 22256342_Written_Assignment.pdf
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import os
import subprocess


PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(PROJECT_ROOT, "22256342_Written_Assignment.docx")
PDF_PATH = os.path.join(PROJECT_ROOT, "22256342_Written_Assignment.pdf")

doc = Document()

# ---- Page margins (A4, 1 inch) ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ---- Default font: Times New Roman 11pt, 1.15 line spacing ----
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)

# ---- Heading styles ----
for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Times New Roman'
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    h_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h_style.paragraph_format.space_before = Pt(10)
    h_style.paragraph_format.space_after = Pt(4)


def add_para(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             size=11, indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_heading(text, level=1, size=None):
    """Custom heading with explicit font control."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.bold = True
    if level == 1:
        r.font.size = Pt(size or 14)
    elif level == 2:
        r.font.size = Pt(size or 12)
    else:
        r.font.size = Pt(size or 11)
    return p


# ============================================================
# Header Block
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("COMP4026 Computer Vision and Pattern Recognition")
tr.font.name = 'Times New Roman'
tr.font.size = Pt(14)
tr.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("Written Assignment (2025-26 Second Semester)")
sr.font.name = 'Times New Roman'
sr.font.size = Pt(12)
sr.bold = True

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Student ID: 22256342")
mr.font.name = 'Times New Roman'
mr.font.size = Pt(11)

# Horizontal separator
sep = doc.add_paragraph()
sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr2 = sep.add_run("—" * 30)
sr2.font.name = 'Times New Roman'
sr2.font.size = Pt(10)

# ============================================================
# PART (i)
# ============================================================
add_heading("Part (i): Selected YOLO-Family Algorithm — YOLOv7", level=1)

add_para(
    "Selected paper: C.-Y. Wang, A. Bochkovskiy, and H.-Y. M. Liao, "
    "\"YOLOv7: Trainable bag-of-freebies sets new state-of-the-art for real-time "
    "object detectors,\" in Proc. IEEE/CVF CVPR, 2023 (arXiv:2207.02696).",
    italic=True, size=10,
)

add_heading("Research Problem", level=2)
add_para(
    "By 2022, the YOLO family (YOLOv4, YOLOv5, YOLOR, YOLOX, PPYOLOE, "
    "Scaled-YOLOv4) had already pushed real-time object detection to a high "
    "level, yet the field was stuck on an unfavourable Pareto frontier: "
    "gaining one or two accuracy points on MS COCO almost always meant "
    "either adding millions of parameters, adding latency at inference, or "
    "adding auxiliary branches that inflate the deployed model. YOLOv7 "
    "explicitly targets this \"speed versus accuracy versus parameter-cost\" "
    "trilemma for general-purpose GPU detectors. More specifically, the "
    "paper identifies three concrete sub-problems. First, standard network "
    "scaling laws (depth-only or width-only scaling in the style of "
    "EfficientNet) are ill-defined for concatenation-based backbones such "
    "as ELAN / CSPVoVNet, because deepening the computational block also "
    "changes the input width of the following transition layer and "
    "corrupts the original gradient-flow geometry. Second, many recent "
    "re-parameterisation tricks (e.g., RepVGG's identity-plus-conv branch) "
    "silently break when they are grafted onto a backbone that already "
    "contains residual or concatenation shortcuts, so they cannot be "
    "ported to YOLO as drop-in upgrades. Third, auxiliary-head "
    "deep-supervision schemes that were known to help training usually "
    "introduce extra inference-time modules or independently-assigned "
    "labels, breaking the \"real-time\" promise of YOLO. The authors frame "
    "YOLOv7 as a search for optimisations that are purely \"trainable\"—they "
    "pay a price only during training and leave the deployed model "
    "untouched in FLOPs and parameters."
)

add_heading("Novelty and Key Contributions", level=2)
add_para(
    "In my own reading, the novelty of YOLOv7 can be organised around four "
    "ideas that together form the so-called trainable bag-of-freebies. "
    "The first is the Extended Efficient Layer Aggregation Network "
    "(E-ELAN). Whereas ELAN aggregates features by controlling the "
    "shortest and longest gradient paths, E-ELAN pushes this idea further "
    "by widening group-convolution cardinality, then shuffling and merging "
    "the resulting feature maps across branches. This lets the block "
    "learn more diverse features while the macro-architecture—and hence "
    "the gradient-path structure that makes the network trainable—remains "
    "unchanged, which is essential for stacking many blocks in deep "
    "variants such as YOLOv7-E6E."
)
add_para(
    "The second idea is compound model scaling tailored to "
    "concatenation-based architectures. Instead of scaling depth and "
    "width independently, YOLOv7 derives a coupled scaling rule that "
    "preserves the original in/out channel ratio of each concatenation "
    "block, so that scaling up the model does not silently distort the "
    "transition layers. Ablation in the paper shows that this compound "
    "rule gives +1.2 AP on COCO, compared to +0.7 AP for width-only and "
    "+1.0 AP for depth-only scaling."
)
add_para(
    "The third idea is planned re-parameterised convolution. The authors "
    "analyse, from a gradient-flow viewpoint, why RepConv's identity "
    "branch damages accuracy when it is placed on top of a ResNet-style "
    "residual or DenseNet-style concatenation connection: the identity "
    "branch duplicates a path that already exists and therefore narrows "
    "the effective gradient diversity. They propose RepConvN, a variant "
    "without the identity connection, and give a simple placement rule "
    "saying that any layer which already has a residual or concatenation "
    "input should use RepConvN instead of RepConv."
)
add_para(
    "The fourth idea is coarse-to-fine label assignment with a lead-head "
    "and an auxiliary head. The lead (final) head generates soft labels "
    "that supervise both itself (fine labels, high precision) and the "
    "auxiliary head (coarse labels, high recall). Because the auxiliary "
    "head is only used during training, this deep-supervision scheme "
    "tightens convergence without adding a single FLOP at inference time. "
    "Combined with minor tricks—batch-norm absorption into convolutions, "
    "YOLOR-style implicit knowledge, and EMA weight averaging—YOLOv7 "
    "reports 56.8 AP on MS COCO at 30+ FPS on V100, surpassing the "
    "transformer-based SWIN-L Cascade-Mask R-CNN by about 509% in speed "
    "and 2 AP, with roughly 40% fewer parameters than comparable YOLO "
    "baselines. Historically, YOLOv7 is important because it is the last "
    "major \"pure-vision\" YOLO milestone before the field pivoted towards "
    "open-vocabulary detection with YOLO-World and YOLOE; most later "
    "real-time detectors still inherit its E-ELAN block and its training-"
    "versus-inference-cost separation as a design philosophy."
)

# Page break before Part (ii)
doc.add_page_break()

# ============================================================
# PART (ii)
# ============================================================
add_heading("Part (ii): Two Unmet Needs of SOTA Visual Object Detection", level=1)
add_para(
    "(Essay — approximately 380 words, under the 400-word limit.)",
    italic=True, size=10,
)

# Essay body
add_para(
    "Modern visual object detectors—from YOLOv7 and YOLOv10 to "
    "Grounding-DINO, YOLO-World and YOLOE—now deliver impressive numbers "
    "on clean MS COCO and even on open-vocabulary LVIS. After reviewing "
    "these methods in class, however, I am convinced that two unmet "
    "needs remain genuinely open, and that neither is closed simply by "
    "\"making the model larger\"."
)

add_para(
    "The first unmet need is robustness under real-world distribution "
    "shift. Benchmarks such as PASCAL-C, COCO-C and Cityscapes-C have "
    "repeatedly shown that detectors trained on clean data lose "
    "30%–60% of their mean Average Precision once the input is corrupted "
    "by fog, rain, snow, motion blur or low-light noise; ICCV 2023's "
    "COCO-O reports a 55.7% relative drop for Faster R-CNN on naturally "
    "shifted images, and CVPR 2023 extends the same story to 3D detectors "
    "on KITTI-C, nuScenes-C and Waymo-C, where camera-only models "
    "essentially collapse under motion-level corruption. What is "
    "important is that neither open-vocabulary alignment (YOLO-World, "
    "Grounding-DINO) nor stronger backbones (ConvNeXt, Swin) close this "
    "gap—they inherit the same fragility from ImageNet-style "
    "pre-training. For safety-critical deployment such as autonomous "
    "driving at night or medical imaging at another hospital, this gap "
    "is not acceptable, yet there is still no detector that is "
    "simultaneously SOTA on COCO and robust across COCO-C/COCO-O."
)

add_para(
    "The second unmet need is a truly favourable accuracy-efficiency "
    "trade-off on edge hardware. The highest accuracy today comes from "
    "transformer detectors such as DINO and Grounding-DINO, whose "
    "backbones routinely exceed 200–500 M parameters and cannot run in "
    "real time on mobile or embedded GPUs. YOLO variants are fast but "
    "give up non-trivial accuracy on complex scenes, and the very "
    "recent YOLO26 (arXiv:2509.25164) benchmark on Jetson Nano and "
    "Jetson Orin confirms that no single model currently achieves SOTA "
    "COCO AP, real-time latency on a Jetson-class device, and a small "
    "enough memory footprint at the same time. Quantisation and "
    "distillation help but recover only part of the lost accuracy. "
    "Until this three-way trade-off is resolved, SOTA detection will "
    "remain a data-centre privilege rather than a property of the "
    "cameras that actually need it."
)

add_para(
    "Closing these two gaps—robust SOTA and edge-deployable SOTA—is, in "
    "my view, the most consequential frontier of visual object detection."
)

# ============================================================
# References
# ============================================================
add_heading("References", level=1)

refs = [
    "[1] C.-Y. Wang, A. Bochkovskiy, and H.-Y. M. Liao, \"YOLOv7: Trainable "
    "bag-of-freebies sets new state-of-the-art for real-time object "
    "detectors,\" in Proc. IEEE/CVF Conf. Computer Vision and Pattern "
    "Recognition (CVPR), 2023. arXiv:2207.02696.",

    "[2] T. Cheng, L. Song, Y. Ge, W. Liu, X. Wang, and Y. Shan, "
    "\"YOLO-World: Real-time open-vocabulary object detection,\" in Proc. "
    "CVPR, 2024. arXiv:2401.17270.",

    "[3] C. Michaelis et al., \"Benchmarking robustness in object "
    "detection: Autonomous driving when winter is coming,\" "
    "arXiv:1907.07484, 2019/2020.",

    "[4] X. Mao et al., \"COCO-O: A benchmark for object detectors under "
    "natural distribution shifts,\" in Proc. IEEE/CVF Int. Conf. Computer "
    "Vision (ICCV), 2023.",

    "[5] Y. Dong et al., \"Benchmarking robustness of 3D object detection "
    "to common corruptions in autonomous driving,\" in Proc. CVPR, 2023. "
    "arXiv:2303.11040.",

    "[6] R. Sapkota, R. H. Cheppally, A. Sharda, and M. Karkee, \"YOLO26: "
    "Key architectural enhancements and performance benchmarking for "
    "real-time object detection,\" arXiv:2509.25164, 2025.",

    "[7] A. Gupta, P. Dollár, and R. Girshick, \"LVIS: A dataset for large "
    "vocabulary instance segmentation,\" in Proc. CVPR, 2019. "
    "arXiv:1908.03195.",
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(r)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)


# ---- Save DOCX ----
doc.save(DOCX_PATH)
print(f"DOCX saved to: {DOCX_PATH}")


# ---- Convert DOCX -> PDF using LibreOffice (macOS) ----
def convert_to_pdf(docx_path, out_dir):
    candidates = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "soffice",
        "libreoffice",
    ]
    for soffice in candidates:
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", out_dir, docx_path],
                check=True, capture_output=True, text=True,
            )
            return True
        except (FileNotFoundError, subprocess.CalledProcessError):
            continue
    return False


if convert_to_pdf(DOCX_PATH, PROJECT_ROOT):
    print(f"PDF saved to: {PDF_PATH}")
else:
    # Fallback: macOS textutil + cupsfilter
    try:
        rtf_path = DOCX_PATH.replace(".docx", ".rtf")
        subprocess.run(
            ["textutil", "-convert", "rtf", DOCX_PATH, "-output", rtf_path],
            check=True,
        )
        subprocess.run(
            ["cupsfilter", rtf_path],
            check=True,
            stdout=open(PDF_PATH, "wb"),
        )
        os.remove(rtf_path)
        print(f"PDF saved (via textutil) to: {PDF_PATH}")
    except Exception as e:
        print(f"WARNING: could not auto-convert to PDF: {e}")
        print("Please open the .docx and export as PDF manually.")
