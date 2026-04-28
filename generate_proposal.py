"""
Generate COMP4026 Project Proposal as a Word Document.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ---- Page margins ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ---- Default font ----
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# ---- Heading styles ----
for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Arial'
    h_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("COMP4026 Computer Vision and Pattern Recognition")
run.font.size = Pt(14)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Group Project Proposal")
run.font.size = Pt(22)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Anonymised Facial Expression Recognition")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)

doc.add_paragraph()
doc.add_paragraph()

info_lines = [
    "Group Members:",
    "[Student A Name] \u2014 Face Recognition Model",
    "[Student B Name] \u2014 Face Anonymisation Model",
    "[Student C Name] \u2014 Expression Recognition Model",
    "",
    "2nd Semester 2025\u20132026",
    "Submission Date: 17 March 2026",
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    if line == "Group Members:":
        run.bold = True

doc.add_page_break()

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Arial'
        # Dark background
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2C3E50',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elm)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Arial'

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacing after table
    return table


# ============================================================
# 1. INTRODUCTION
# ============================================================
doc.add_heading('1. Introduction', level=1)

doc.add_paragraph(
    "Facial recognition and expression recognition systems have achieved remarkable "
    "accuracy in recent years. However, these systems process biometric data that contains "
    "sensitive personal information, raising significant privacy concerns. The European "
    "Union\u2019s GDPR, for example, classifies facial images as biometric data requiring "
    "explicit consent for processing. This creates a fundamental tension: how can we "
    "leverage the utility of facial expression analysis while protecting individual identity?"
)

p = doc.add_paragraph()
p.add_run("This project addresses this challenge by building a prototype for ")
p.add_run("Anonymised Facial Expression Recognition").bold = True
p.add_run(
    ". The core idea is to transform face images into anonymised versions that preserve "
    "facial expressions while concealing the original identity. The system consists of "
    "three interconnected components: (1) a face recognition model that verifies identity "
    "protection strength, (2) a face anonymisation model that transforms images, and "
    "(3) a facial expression recognition model that works on the anonymised output."
)

doc.add_paragraph(
    "The key challenge is to maintain expression recognition accuracy while ensuring "
    "that a well-trained face recognition model cannot identify the original person from "
    "the anonymised images. A successful system would achieve near-random-chance identity "
    "recognition on anonymised images while preserving high expression classification accuracy."
)

# ============================================================
# 2. PRELIMINARY STUDIES
# ============================================================
doc.add_heading('2. Preliminary Studies', level=1)

doc.add_heading('2.1 Dataset Analysis', level=2)
p = doc.add_paragraph()
p.add_run(
    "We have conducted preliminary analysis on the assigned datasets. For the face "
    "recognition component (Student A), we use the "
)
p.add_run("Pins Face Recognition").bold = True
p.add_run(
    " dataset from Kaggle (hereisburak/pins-face-recognition), which contains "
    "17,534 face images across 105 celebrity identities. The dataset exhibits moderate "
    "class imbalance, with images per identity ranging from 86 (Lionel Messi) to 237 "
    "(Leonardo DiCaprio), and a mean of approximately 167 images per identity."
)

p = doc.add_paragraph()
p.add_run("Table 1: Dataset Overview").bold = True
p.add_run().font.size = Pt(10)
add_table(doc,
    ["Dataset", "Purpose", "Details"],
    [
        ["Pins Face Recognition", "Face Recognition (Student A)", "105 identities, 17,534 images"],
        ["CelebA-HQ", "Face Anonymisation (Student B)", "High-quality celebrity faces"],
        ["FER-2013", "Expression Classification (Student C)", "7 expressions, 35,887 images"],
    ],
    col_widths=[5, 5.5, 5.5],
)

doc.add_paragraph(
    "Our initial data exploration reveals that the Pins Face Recognition dataset consists "
    "of internet-sourced images with diverse lighting conditions, poses, and backgrounds, "
    "making it suitable for training a robust face recognition model. The data split is "
    "70% training (12,225 images), 15% validation (2,574 images), and 15% testing "
    "(2,735 images)."
)

doc.add_heading('2.2 Preliminary Experiments', level=2)
doc.add_paragraph(
    "We conducted a quick-test preliminary experiment using a subset of 20 identities "
    "from the Pins dataset to validate our face recognition pipeline. Using a ResNet-50 "
    "backbone with ArcFace loss, trained for only 5 epochs on an Apple M-series GPU (MPS), "
    "we achieved a test accuracy of 90.65% (Top-1) and 99.39% (Top-5). This confirms "
    "that the chosen architecture and dataset are viable for building a reliable identity "
    "classifier."
)

p = doc.add_paragraph()
p.add_run("Table 2: Preliminary Experiment Results (20-Identity Subset)").bold = True
add_table(doc,
    ["Metric", "Value"],
    [
        ["Top-1 Identification Accuracy", "90.65%"],
        ["Top-5 Identification Accuracy", "99.39%"],
        ["Equal Error Rate (EER)", "0.069"],
        ["TAR @ FAR=0.01", "0.792"],
        ["Training Time (5 epochs, MPS)", "7.3 minutes"],
    ],
    col_widths=[8, 8],
)

doc.add_paragraph(
    "These preliminary results on 20 identities validated the pipeline and informed "
    "our two-phase training strategy (see Section 3). We then proceeded to full-scale "
    "training on all 105 identities."
)

doc.add_heading('2.3 Full Training Results (105 Identities)', level=2)
doc.add_paragraph(
    "Following the preliminary experiments, we trained the face recognition model on "
    "all 105 identities in the Pins Face Recognition dataset for 30 epochs using our "
    "two-phase training strategy: Cross-Entropy warmup for the first 5 epochs followed "
    "by ArcFace loss (scale=32, margin=0.3) for the remaining 25 epochs. Training was "
    "performed on an Apple M-series GPU (MPS) and took approximately 291.8 minutes."
)

p = doc.add_paragraph()
p.add_run("Table 3: Full Training Results (105 Identities, 30 Epochs)").bold = True
add_table(doc,
    ["Metric", "Value"],
    [
        ["Top-1 Identification Accuracy", "93.46%"],
        ["Top-5 Identification Accuracy", "98.03%"],
        ["Best Validation Accuracy", "93.01%"],
        ["Equal Error Rate (EER)", "0.0528"],
        ["TAR @ FAR=0.1", "0.9694"],
        ["TAR @ FAR=0.01", "0.8576"],
        ["TAR @ FAR=0.001", "0.7326"],
        ["Test Loss", "1.699"],
        ["Training Time (30 epochs, MPS)", "291.8 minutes"],
    ],
    col_widths=[8, 8],
)

doc.add_paragraph(
    "The two-phase training strategy proved essential for achieving strong performance. "
    "During the CE warmup phase (epochs 1\u20135), validation accuracy reached approximately "
    "80%. After switching to ArcFace loss at epoch 6, the model steadily improved to a "
    "final test accuracy of 93.46%. The training curves show healthy convergence with "
    "a clear phase transition at epoch 6. The confusion matrix reveals a strong diagonal "
    "pattern, indicating consistent correct predictions across all 105 identities."
)

doc.add_paragraph(
    "The verification metrics further confirm the model\u2019s discriminative capability. "
    "An EER of 0.0528 means that at the equal error operating point, only 5.28% of "
    "genuine pairs are rejected and 5.28% of impostor pairs are incorrectly accepted. "
    "The TAR@FAR=0.01 of 0.8576 indicates that 85.76% of genuine pairs are correctly "
    "verified when the false acceptance rate is held at 1%. These results establish a "
    "strong baseline for evaluating the anonymisation model\u2019s identity concealment strength."
)

# ============================================================
# 3. METHODOLOGY
# ============================================================
doc.add_heading('3. Methodology', level=1)

doc.add_heading('3.1 System Architecture', level=2)
doc.add_paragraph(
    "The overall system follows a three-stage pipeline: (1) the anonymisation model "
    "transforms an input face image to remove identity while preserving expression, "
    "(2) the expression recognition model classifies the emotion in the anonymised image, "
    "and (3) the face recognition model evaluates whether the original identity can still "
    "be recovered. The face recognition model serves as both a standalone classifier and "
    "a privacy evaluation tool."
)

doc.add_heading('3.2 Face Recognition Model (Student A)', level=2)

p = doc.add_paragraph()
p.add_run("Architecture: ").bold = True
p.add_run(
    "The face recognition model uses a ResNet-50 backbone pretrained on ImageNet, followed "
    "by a fully connected embedding layer (2048 \u2192 512 dimensions) with Batch Normalisation. "
    "The output is an L2-normalised 512-dimensional face embedding. During training, an "
    "ArcFace classification head applies additive angular margin to the cosine similarity "
    "between embeddings and class prototypes, producing highly discriminative features."
)

p = doc.add_paragraph()
p.add_run("Training Strategy: ").bold = True
p.add_run(
    "Based on our preliminary experiments, we adopt a two-phase training approach. "
    "Phase 1 (Warmup, epochs 1\u20135) uses standard Cross-Entropy loss to establish a good "
    "initial embedding space. Phase 2 (Fine-tuning, epochs 6\u201330) switches to ArcFace loss "
    "(scale=32, margin=0.3) for discriminative feature learning. We use differential "
    "learning rates: 0.0001 for the pretrained backbone and 0.001 for the new embedding "
    "and classification heads. AdamW optimiser with weight decay of 5e-4 and cosine "
    "learning rate scheduling are employed."
)

p = doc.add_paragraph()
p.add_run("Justification: ").bold = True
p.add_run(
    "We chose ResNet-50 for its strong representation learning capability and wide adoption "
    "in face recognition research [2]. ArcFace loss was selected because it directly optimises "
    "the angular discriminability of embeddings, which is superior to softmax Cross-Entropy "
    "for face recognition tasks [1]. The two-phase training strategy addresses the convergence "
    "difficulty that arises when ArcFace margin is applied from the beginning to a randomly "
    "initialised head. Label smoothing (0.1) is used to prevent overconfident predictions "
    "and improve generalisation."
)

p = doc.add_paragraph()
p.add_run("Table 4: Face Recognition Model Configuration").bold = True
add_table(doc,
    ["Component", "Configuration"],
    [
        ["Backbone", "ResNet-50 (ImageNet V2 pretrained)"],
        ["Embedding Dimension", "512"],
        ["Loss Function", "ArcFace (s=32, m=0.3) with CE warmup"],
        ["Input Size", "160 \u00d7 160 pixels"],
        ["Optimiser", "AdamW (lr=1e-3 head, 1e-4 backbone)"],
        ["Total Parameters", "24.6 million"],
        ["Training Epochs", "30 (5 CE warmup + 25 ArcFace)"],
        ["Data Augmentation", "Random flip, rotation (\u00b110\u00b0), colour jitter"],
    ],
    col_widths=[8, 8],
)

doc.add_heading('3.3 Face Anonymisation Model (Student B)', level=2)
doc.add_paragraph(
    "The anonymisation model will transform face images to conceal identity while preserving "
    "expression. Potential approaches include GAN-based face de-identification (e.g., "
    "DeepPrivacy [4], CIAGAN [6]), diffusion-based methods, or k-Same-Net approaches. "
    "The model will be trained on CelebA-HQ and evaluated by measuring how much identity "
    "information is removed (using Student A\u2019s model) while retaining expression information "
    "(using Student C\u2019s model)."
)

doc.add_heading('3.4 Expression Recognition Model (Student C)', level=2)
doc.add_paragraph(
    "The expression recognition model will classify facial expressions into seven standard "
    "categories (angry, disgust, fear, happy, sad, surprise, neutral) using the FER-2013 "
    "dataset [5]. A CNN-based architecture will be used, potentially leveraging transfer "
    "learning from ImageNet. The model must maintain high accuracy when applied to anonymised "
    "images, which may exhibit visual artefacts or altered appearance from the anonymisation process."
)

# ============================================================
# 4. EVALUATION PLAN
# ============================================================
doc.add_heading('4. Evaluation Plan', level=1)
doc.add_paragraph(
    "The evaluation framework measures two competing objectives: privacy protection "
    "(identity concealment) and utility preservation (expression accuracy)."
)

p = doc.add_paragraph()
p.add_run("Table 5: Evaluation Metrics").bold = True
add_table(doc,
    ["Category", "Metric", "Target"],
    [
        ["Privacy", "ID accuracy on original images", ">90% (achieved: 93.46%)"],
        ["Privacy", "ID accuracy on anonymised images", "Near random chance (~1%)"],
        ["Privacy", "EER (Equal Error Rate)", "High EER on anonymised"],
        ["Utility", "Expression consistency rate", ">80%"],
        ["Utility", "Expression accuracy on anonymised", "Close to original accuracy"],
    ],
    col_widths=[3, 6, 7],
)

doc.add_paragraph(
    "The privacy protection rate is defined as (100% \u2212 anonymised recognition accuracy). "
    "An ideal anonymisation model would reduce the face recognition accuracy from >90% to "
    "near random chance (approximately 0.95% for 105 identities), indicating that the "
    "original identity is effectively concealed."
)

# ============================================================
# 5. INITIAL DESIGN
# ============================================================
doc.add_heading('5. Initial Design and Implementation Plan', level=1)
doc.add_paragraph(
    "The project follows a modular design where each component can be developed independently "
    "and later integrated for end-to-end evaluation. Student A\u2019s face recognition model has "
    "been fully implemented and trained on all 105 identities, achieving 93.46% Top-1 "
    "identification accuracy and an EER of 0.0528. The implementation uses PyTorch and "
    "provides a clean API (FaceRecognizer class) for other team members to evaluate "
    "anonymisation quality programmatically. The model is now ready for integration with "
    "the anonymisation pipeline."
)

p = doc.add_paragraph()
p.add_run("Table 6: Implementation Timeline").bold = True
add_table(doc,
    ["Period", "Task", "Responsible"],
    [
        ["Week 1\u20133", "Individual model development and preliminary training", "A, B, C"],
        ["Week 4\u20135", "Model optimisation and hyperparameter tuning", "A, B, C"],
        ["Week 6", "Integration: connect anonymisation output to recognition models", "A, B, C"],
        ["Week 7", "End-to-end evaluation and comparative analysis", "A, B, C"],
        ["Week 8", "Presentation preparation and report writing", "A, B, C"],
    ],
    col_widths=[3, 10, 3],
)

# ============================================================
# 6. REFERENCES
# ============================================================
doc.add_heading('6. References', level=1)
refs = [
    "[1] J. Deng, J. Guo, N. Xue, and S. Zafeiriou, \u201cArcFace: Additive Angular Margin Loss for Deep Face Recognition,\u201d in Proc. CVPR, 2019.",
    "[2] K. He, X. Zhang, S. Ren, and J. Sun, \u201cDeep Residual Learning for Image Recognition,\u201d in Proc. CVPR, 2016.",
    "[3] F. Schroff, D. Kalenichenko, and J. Philbin, \u201cFaceNet: A Unified Embedding for Face Recognition and Clustering,\u201d in Proc. CVPR, 2015.",
    "[4] H. Hukkel\u00e5s, R. Mester, and F. Lindseth, \u201cDeepPrivacy: A Generative Adversarial Network for Face Anonymization,\u201d in Proc. ISVC, 2019.",
    "[5] I. J. Goodfellow et al., \u201cChallenges in Representation Learning: A Report on Three Machine Learning Contests,\u201d in Neural Networks, vol. 64, pp. 59\u201363, 2015.",
    "[6] M. Maximov, I. Elezi, and L. Leal-Taix\u00e9, \u201cCIAGAN: Conditional Identity Anonymization Generative Adversarial Networks,\u201d in Proc. CVPR, 2020.",
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)

# ============================================================
# 7. GenAI DISCLOSURE
# ============================================================
doc.add_heading('7. GenAI Usage Disclosure', level=1)
doc.add_paragraph(
    "Generative AI tools were used in the following aspects of this project:"
)

disclosures = [
    ("Code scaffolding: ", "AI assistance was used to generate initial boilerplate code for the training pipeline, data loading utilities, and evaluation scripts. All generated code was reviewed, tested, and modified by team members."),
    ("Documentation: ", "AI tools assisted in drafting documentation and this proposal. Content was reviewed and edited for accuracy."),
    ("Architecture design: ", "The choice of ResNet-50 + ArcFace architecture was informed by academic literature and validated through our own preliminary experiments."),
]
for bold_text, normal_text in disclosures:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(bold_text).bold = True
    p.add_run(normal_text)

doc.add_paragraph(
    "All experimental results, analysis, and conclusions in this proposal are based on "
    "actual code execution and real experimental outcomes."
)

# ---- Save ----
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Project_Proposal.docx")
doc.save(output_path)
print(f"Proposal generated: {output_path}")
